from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Márgenes ────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

# ─── Paleta de colores ───────────────────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x1A, 0x37, 0x6C)   # títulos principales
AZUL_MEDIO   = RGBColor(0x21, 0x5C, 0xA0)   # subtítulos de fase
AZUL_CLARO   = RGBColor(0xDE, 0xEA, 0xF6)   # fondo cabeceras tabla
VERDE_PISTA  = RGBColor(0x1E, 0x6B, 0x42)   # texto bloque pistas
FONDO_PISTA  = RGBColor(0xE8, 0xF5, 0xEC)   # fondo bloque pistas
NARANJA      = RGBColor(0xC5, 0x5A, 0x11)   # entregable
GRIS_LINEA   = RGBColor(0xCC, 0xCC, 0xCC)

# ─── Helpers de estilo ────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name  = "Calibri"
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=20, bold=True, color=AZUL_OSCURO)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=15, bold=True, color=AZUL_OSCURO)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    # línea inferior
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A376C')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_heading3(text, color=AZUL_MEDIO):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=color)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_heading_fase(numero, titulo, estimacion=None):
    p = doc.add_paragraph()
    r1 = p.add_run(f"FASE {numero}  ")
    set_font(r1, size=13, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
    r2 = p.add_run(titulo)
    set_font(r2, size=13, bold=True, color=AZUL_OSCURO)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    # línea de separación superior
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '215CA0')
    pBdr.append(top)
    pPr.append(pBdr)
    return p

def add_body(text, bold_parts=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=11)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=11)
    p.paragraph_format.left_indent = Cm(1 + level * 0.8)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_numbered(text, level=0):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    set_font(run, size=11)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_pista(lines):
    """Bloque verde de pistas."""
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        prefix = "💡 " if i == 0 else "    "
        run = p.add_run(prefix + line)
        set_font(run, size=10, italic=(i > 0), color=VERDE_PISTA)
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.right_indent = Cm(0.8)
        p.paragraph_format.space_after  = Pt(1)
        # fondo verde claro
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'E8F5EC')
        pPr.append(shd)

def add_entregable(text):
    p = doc.add_paragraph()
    r1 = p.add_run("✅ Entregable: ")
    set_font(r1, size=11, bold=True, color=NARANJA)
    r2 = p.add_run(text)
    set_font(r2, size=11, italic=True, color=NARANJA)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

def add_code_inline(text):
    """Párrafo de código (monoespaciado, fondo gris)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)

def add_table_simple(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # cabecera
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        set_font(run, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '215CA0')
        tcPr.append(shd)
    # filas
    for r_idx, row_data in enumerate(rows):
        row = t.rows[r_idx + 1]
        fill = 'FFFFFF' if r_idx % 2 == 0 else 'EEF3FA'
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            set_font(run, size=10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PRÁCTICA FINAL")
set_font(run, size=28, bold=True, color=AZUL_OSCURO)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Limpieza y normalización de datos con Python")
set_font(run, size=16, italic=True, color=AZUL_MEDIO)

doc.add_paragraph()

# Caja de metadatos
t = doc.add_table(rows=4, cols=2)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
meta = [
    ("Modalidad",     "Individual"),
    ("Entrega",       "Al finalizar las sesiones de clase"),
    ("Herramientas",  "Python 3.12 y los recursos vistos en clase"),
    ("Restricción",   "❌ No se pueden usar librerías no vistas en clase"),
]
for i, (k, v) in enumerate(meta):
    row = t.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    run0 = row.cells[0].paragraphs[0].runs[0]
    run1 = row.cells[1].paragraphs[0].runs[0]
    set_font(run0, size=11, bold=True, color=AZUL_OSCURO)
    set_font(run1, size=11)
    fill0 = 'DEE8F4'
    for cell, fill in [(row.cells[0], fill0), (row.cells[1], 'FFFFFF')]:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  INTRODUCCIÓN
# ══════════════════════════════════════════════════════════════════════════════

add_heading1("¿EN QUÉ CONSISTE ESTA PRÁCTICA?")

add_body(
    "En el mundo real, los datos casi nunca llegan limpios. Vienen de sistemas distintos, "
    "los introducen personas diferentes y acaban llenos de inconsistencias: nombres escritos "
    "de diez formas distintas, fechas en formatos imposibles, valores vacíos, duplicados y "
    "cifras que no tienen ningún sentido."
)
add_body(
    "Tu trabajo en esta práctica es exactamente ese: recibir datos sucios, entender qué "
    "falla, limpiarlos con Python y dejarlos listos para usar. Sin pandas. Sin atajos. "
    "Solo Python puro y las herramientas que ya conoces."
)
add_body(
    "Tienes dos opciones temáticas. El trabajo técnico es equivalente en ambas. Elige la "
    "que más te motive."
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  OPCIÓN A
# ══════════════════════════════════════════════════════════════════════════════

add_heading2("OPCIÓN A — FESTIVAL \"SOUNDWAVE 2026\"")

add_heading3("La historia", color=AZUL_MEDIO)
add_body(
    "Te acaban de contratar como técnico de datos en SoundWave, uno de los festivales de "
    "música más grandes de España. Faltan tres semanas para que arranque el festival y todo "
    "es un caos."
)
add_body(
    "El equipo de producción, el de ventas, el de marketing y los managers de los artistas "
    "han ido volcando información en ficheros separados, cada uno a su manera. Hay datos "
    "duplicados, nombres escritos de formas distintas, fechas en formatos imposibles, campos "
    "vacíos por todas partes y cifras que no cuadran."
)

p = doc.add_paragraph()
r1 = p.add_run("Tu jefe Carlos te lo deja claro: ")
set_font(r1, size=11)
r2 = p.add_run(
    "\"Necesito que TODOS los datos estén limpios, unificados y listos para cargar en el "
    "sistema del festival antes de la próxima reunión. Si subimos estos datos tal cual, "
    "el festival será un desastre.\""
)
set_font(r2, size=11, italic=True, color=AZUL_OSCURO)

add_heading3("Los ficheros  (carpeta opcion_A_festival/datos/)", color=AZUL_MEDIO)

add_table_simple(
    ["Fichero", "Contenido", "Registros aprox."],
    [
        ["artistas.csv",             "Artistas contratados. Caché, país, contacto manager.", "~250 filas"],
        ["escenarios_horarios.xlsx", "Programación: qué artista toca en qué escenario y cuándo.", "~200 filas"],
        ["ventas_entradas.json",     "Registro de venta de entradas con comprador, tipo y precio.", "~400 entradas"],
        ["patrocinadores.xml",       "Empresas patrocinadoras: importe, categoría y fechas de contrato.", "~40 entradas"],
    ],
    col_widths=[5, 8, 3.5]
)

# ══════════════════════════════════════════════════════════════════════════════
#  OPCIÓN B
# ══════════════════════════════════════════════════════════════════════════════

add_heading2("OPCIÓN B — LIGA DE ESPORTS \"CYBERLEAGUE PRO\"")

add_heading3("La historia", color=AZUL_MEDIO)
add_body(
    "La CyberLeague Pro es una de las ligas de eSports con más crecimiento del momento. "
    "Tras dos temporadas, la liga necesita publicar las estadísticas oficiales en su web "
    "y compartir los datos con patrocinadores e inversores."
)
add_body(
    "Los datos los han ido recogiendo diferentes personas: árbitros, managers, el equipo "
    "de comunicación y el departamento financiero, cada uno con su propio criterio. "
    "El resultado: jugadores duplicados, equipos con nombres distintos en cada fichero, "
    "partidas con resultados imposibles y premios que no cuadran."
)

p = doc.add_paragraph()
r1 = p.add_run("Laura, directora de operaciones, te ha contratado: ")
set_font(r1, size=11)
r2 = p.add_run(
    "\"Nos jugamos la credibilidad de la liga. Si publicamos estos datos así, los equipos "
    "van a protestar, los patrocinadores van a desconfiar y la prensa nos va a crucificar. "
    "Necesito que todo esté impecable.\""
)
set_font(r2, size=11, italic=True, color=AZUL_OSCURO)

add_heading3("Los ficheros  (carpeta opcion_B_esports/datos/)", color=AZUL_MEDIO)

add_table_simple(
    ["Fichero", "Contenido", "Registros aprox."],
    [
        ["jugadores.csv",  "Jugadores de la liga: gamertag, edad, equipo, rol y salario.", "~250 filas"],
        ["equipos.xlsx",   "Equipos (hoja 1) y staff técnico (hoja 2).", "~90 filas en 2 hojas"],
        ["partidas.json",  "Resultados de partidas: equipos, mapa, puntuación y torneo.", "~400 entradas"],
        ["premios.csv",    "Premios de torneos: equipo, posición y cuantía económica.", "~150 filas"],
    ],
    col_widths=[5, 8, 3.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  ESTRUCTURA DE ENTREGA
# ══════════════════════════════════════════════════════════════════════════════

add_heading2("ESTRUCTURA DE ENTREGA")

add_body("Tu carpeta de entrega debe tener exactamente esta estructura:")

for line in [
    "opcion_X/",
    "├── datos/                     ← Ficheros originales (NO los modifiques)",
    "│   └── (ficheros originales tal como los recibiste)",
    "│",
    "├── datos_limpios/             ← Generados por tu código",
    "│   ├── artistas_limpio.csv    ← Un CSV por cada fichero original",
    "│   ├── ...",
    "│   ├── datos_completos.xlsx   ← Excel con todos los datos en hojas separadas",
    "│   └── informe_limpieza.txt   ← Informe del proceso completo",
    "│",
    "├── lib/                       ← Módulos con las funciones organizadas por responsabilidad",
    "│   ├── carga.py               ← Funciones para cargar CSV, Excel, JSON y XML",
    "│   ├── limpieza.py            ← Funciones de limpieza de texto, números, fechas y categorías",
    "│   ├── auditoria.py           ← Funciones de análisis y detección de problemas",
    "│   └── exportacion.py         ← Funciones para generar los ficheros de salida",
    "│",
    "└── main.py                    ← Flujo principal: importa los módulos y orquesta el proceso",
]:
    add_code_inline(line)

add_body(
    "La carpeta lib/ separa el código por responsabilidad: cada fichero agrupa las funciones "
    "relacionadas con una misma tarea. El fichero main.py únicamente importa esos módulos y "
    "ejecuta el flujo completo. Este patrón es el mismo que hemos usado en los ejercicios del curso."
)
add_pista([
    "Cómo importar tus módulos desde main.py:",
    "from lib.carga      import cargar_csv, cargar_excel, cargar_json, cargar_xml",
    "from lib.auditoria  import auditar_fichero",
    "from lib.limpieza   import limpiar_texto, normalizar_fecha, normalizar_categoria",
    "from lib.exportacion import exportar_csv, exportar_excel, generar_informe",
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  FASES
# ══════════════════════════════════════════════════════════════════════════════

add_heading1("LAS 10 FASES DEL PROYECTO")

add_body(
    "Sigue las fases en orden. Cada una produce un resultado visible que puedes verificar "
    "antes de pasar a la siguiente. No avances si la fase anterior no funciona correctamente."
)

# ─── FASE 1 ───────────────────────────────────────────────────────────────────
add_heading_fase(1, "EXPLORACIÓN Y CARGA DE DATOS", "⏱ ~2 horas")

add_heading3("Objetivo")
add_body("Cargar los 4 ficheros en memoria como listas de diccionarios y entender qué contiene cada uno.")

add_heading3("Tareas")
add_numbered("Crea la estructura de carpetas del proyecto: main.py en la raíz y la carpeta lib/ con los ficheros carga.py, limpieza.py, auditoria.py y exportacion.py.")
add_numbered("Crea una función de carga para cada formato en lib/carga.py:")
add_bullet("cargar_csv(ruta) → lista de diccionarios", level=1)
add_bullet("cargar_excel(ruta) → lista de diccionarios  (usa openpyxl)", level=1)
add_bullet("cargar_json(ruta) → lista de diccionarios", level=1)
add_bullet("cargar_xml(ruta) → lista de diccionarios  (solo Opción A)", level=1)
add_body("Cada función recibe la ruta del fichero y devuelve una lista de diccionarios.")
add_numbered("Carga los 4 ficheros y guárdalos en 4 variables.")
add_numbered("Para cada fichero cargado, imprime por pantalla:")
add_bullet("Nombre del fichero", level=1)
add_bullet("Número total de registros", level=1)
add_bullet("Nombres de los campos/columnas", level=1)
add_bullet("Los 5 primeros registros", level=1)

add_pista([
    "Pistas:",
    "Para CSV usa el módulo csv con csv.DictReader — devuelve cada fila como un diccionario directamente.",
    "Para Excel, recuerda iterar desde min_row=2 para saltar la cabecera y usa dict(zip(cabeceras, fila)).",
    "Para JSON, json.load() ya devuelve una lista de diccionarios si el fichero tiene ese formato.",
    "Para XML (solo Opción A), recorre los elementos hijo con root.findall() y construye un diccionario con {campo: elemento.find(campo).text}.",
    "El Excel de Opción B tiene 2 hojas — tendrás que cargar cada una por separado.",
])

add_entregable("Funciones de carga funcionando + salida por pantalla con el resumen de cada fichero.")

# ─── FASE 2 ───────────────────────────────────────────────────────────────────
add_heading_fase(2, "AUDITORÍA DE CALIDAD", "⏱ ~2.5 horas")

add_heading3("Objetivo")
add_body("Analizar cada fichero y detectar todos los problemas que contiene. Antes de limpiar, necesitas saber exactamente qué está mal.")

add_heading3("Tareas")
add_numbered("Para cada fichero, crea una función que analice y cuente:")
add_bullet("Valores vacíos por campo: None, cadena vacía, \"N/A\", \"-\", \"no disponible\", etc.", level=1)
add_bullet("Duplicados exactos y parciales (mismo dato escrito diferente, ej: \"Los Rebeldes\" vs \"los rebeldes\").", level=1)
add_bullet("Formatos inconsistentes: ¿cuántas variaciones distintas hay para un mismo tipo de campo?", level=1)
add_bullet("Valores fuera de rango: precios negativos, edades imposibles, etc.", level=1)
add_bullet("Espacios extra: campos con espacios al principio, al final o dobles en medio.", level=1)

add_numbered("Almacena los resultados en un diccionario de auditoría con esta estructura (o similar):")
add_code_inline("auditoria = {")
add_code_inline("    \"nombre_fichero\": {")
add_code_inline("        \"total_registros\": ...,")
add_code_inline("        \"valores_vacios\": {\"campo1\": X, \"campo2\": Y},")
add_code_inline("        \"duplicados\": X,")
add_code_inline("        \"formatos_inconsistentes\": {\"campo\": [\"var1\", \"var2\", ...]},")
add_code_inline("        \"fuera_de_rango\": {\"campo\": X},")
add_code_inline("        \"espacios_extra\": {\"campo\": X}")
add_code_inline("    }")
add_code_inline("}")

add_numbered("Imprime un informe de auditoría claro y legible para cada fichero. Ejemplo:")
add_code_inline("=== AUDITORÍA: artistas.csv ===")
add_code_inline("Total registros: 250")
add_code_inline("Valores vacíos:")
add_code_inline("  - email_manager: 32 vacíos")
add_code_inline("Duplicados: 15 registros duplicados")
add_code_inline("Formatos inconsistentes:")
add_code_inline("  - genero_musical: 23 variaciones (Rock, rock, ROCK, Rokc...)")
add_code_inline("Espacios extra:")
add_code_inline("  - nombre: 45 campos con espacios extra")

add_pista([
    "Pistas:",
    "Para detectar vacíos, crea una lista de valores que consideras 'vacío': [\"\", None, \"N/A\", \"-\", \"no disponible\"].",
    "Para contar variaciones de formato usa un set(): añade cada valor en minúsculas y cuenta cuántos valores únicos hay.",
    "Para detectar duplicados parciales, normaliza la clave (lower + strip) antes de comparar.",
])

add_entregable("Funciones de auditoría y el informe impreso para los 4 ficheros.")

# ─── FASE 3 ───────────────────────────────────────────────────────────────────
add_heading_fase(3, "LIMPIEZA DE TEXTO Y ESPACIOS", "⏱ ~1.5 horas")

add_heading3("Objetivo")
add_body("Corregir todos los problemas relacionados con formato de texto: espacios extra, mayúsculas inconsistentes y tildes que faltan.")

add_heading3("Tareas")
add_numbered("Crea una función limpiar_texto(texto) que:")
add_bullet("Elimine espacios al principio y al final (.strip()).", level=1)
add_bullet("Elimine espacios dobles o múltiples en medio del texto.", level=1)
add_bullet("Devuelva el texto limpio sin cambiar mayúsculas/minúsculas.", level=1)

add_numbered("Crea una función normalizar_texto(texto) que:")
add_bullet("Aplique limpiar_texto primero.", level=1)
add_bullet("Convierta el texto a formato estándar (primera letra en mayúscula, resto en minúsculas).", level=1)
add_bullet("Corrija tildes que faltan usando el diccionario de correcciones que se proporciona más abajo.", level=1)

add_numbered("Aplica estas funciones a TODOS los campos de texto de los 4 ficheros.")
add_numbered("Imprime cuántas correcciones se han hecho en cada fichero y campo:")
add_code_inline("artistas.csv > nombre: 45 textos corregidos")

add_heading3("Diccionario de correcciones de tildes (punto de partida)")
add_body("Se te proporciona este diccionario como base. Si encuentras más casos en los datos que no estén aquí, añádelos tú.")
add_code_inline("correcciones_tildes = {")
add_code_inline("    \"jose\":       \"José\",      \"maria\":      \"María\",")
add_code_inline("    \"garcia\":     \"García\",    \"gonzalez\":   \"González\",")
add_code_inline("    \"martinez\":   \"Martínez\",  \"lopez\":      \"López\",")
add_code_inline("    \"perez\":      \"Pérez\",     \"sanchez\":    \"Sánchez\",")
add_code_inline("    \"gomez\":      \"Gómez\",     \"fernandez\":  \"Fernández\",")
add_code_inline("    \"rodriguez\":  \"Rodríguez\", \"hernandez\":  \"Hernández\",")
add_code_inline("    \"ramirez\":    \"Ramírez\",   \"gutierrez\":  \"Gutiérrez\",")
add_code_inline("}")

add_pista([
    "Pistas:",
    "Para eliminar espacios dobles: usa un bucle while '  ' in texto: texto = texto.replace('  ', ' ')",
    "  o más directo: ' '.join(texto.split())",
    "Para las tildes: divide el texto en palabras con split(), aplica el diccionario palabra a palabra,",
    "  y vuelve a unir con join(). Así evitas corregir partes de palabras más largas.",
    "Ejemplo: 'Juan Garcia Lopez' → split() → ['Juan','Garcia','Lopez']",
    "  → corrección → ['Juan','García','López'] → join() → 'Juan García López'",
])

add_entregable("Funciones de limpieza y los datos con textos corregidos.")

# ─── FASE 4 ───────────────────────────────────────────────────────────────────
add_heading_fase(4, "LIMPIEZA DE VALORES VACÍOS Y TIPOS", "⏱ ~2 horas")

add_heading3("Objetivo")
add_body("Tratar los valores vacíos y asegurar que cada campo tiene el tipo de dato correcto.")

add_heading3("Tareas")
add_numbered("Define las reglas para cada campo de cada fichero. Para cada campo, decide:")
add_bullet("¿Qué tipo de dato debe tener? (texto, entero, decimal...)", level=1)
add_bullet("Si está vacío: ¿rellenar con un valor por defecto? ¿marcarlo como \"SIN DATOS\"? ¿eliminar el registro?", level=1)
add_body("Documenta estas reglas como comentarios en tu código o en un diccionario.")

add_numbered("Crea una función limpiar_valor_numerico(valor) que:")
add_bullet("Reciba un valor que debería ser numérico pero puede venir como texto con símbolos (\"15.000€\", \"15,000\", \"$500\").", level=1)
add_bullet("Elimine símbolos de moneda (€, $) y separadores de miles.", level=1)
add_bullet("Convierta comas decimales a puntos si es necesario.", level=1)
add_bullet("Devuelva el número como int o float según corresponda.", level=1)
add_bullet("Si no se puede convertir, devuelva None.", level=1)

add_numbered("Aplica la conversión de tipos a todos los campos numéricos.")
add_numbered("Aplica el tratamiento de valores vacíos según las reglas definidas.")
add_numbered("Imprime un resumen de las acciones realizadas:")
add_code_inline("artistas.csv > cache_eur: 5 valores convertidos a número")
add_code_inline("artistas.csv > email_manager: 32 vacíos → marcados como 'SIN DATOS'")

add_pista([
    "Pistas:",
    "Para limpiar un número: str(valor).strip().replace('€','').replace('$','').replace('.','').replace(',','.')",
    "  Ojo con el orden: elimina el punto de miles ANTES de convertir la coma decimal a punto.",
    "Usa try/except ValueError cuando hagas float() o int() para capturar valores que no son convertibles.",
    "Un campo es 'crítico' si sin él el registro no tiene sentido (ej: el nombre de un artista).",
    "  En ese caso, eliminar el registro puede ser la decisión correcta. Documéntalo.",
])

add_entregable("Datos con tipos correctos y valores vacíos tratados.")

# ─── FASE 5 ───────────────────────────────────────────────────────────────────
add_heading_fase(5, "NORMALIZACIÓN DE CATEGORÍAS", "⏱ ~2 horas")

add_heading3("Objetivo")
add_body("Unificar todos los valores que representan lo mismo pero están escritos de formas diferentes.")

add_heading3("Tareas")
add_numbered("Identifica todos los campos que son 'categorías' (valores de un conjunto limitado).")
p = doc.add_paragraph()
r = p.add_run("Opción A: ")
set_font(r, size=11, bold=True)
r2 = p.add_run("genero_musical, pais, escenario, tipo_entrada, metodo_pago, categoria (patrocinador).")
set_font(r2, size=11)
p = doc.add_paragraph()
r = p.add_run("Opción B: ")
set_font(r, size=11, bold=True)
r2 = p.add_run("pais, equipo, rol, region, cargo, mapa, torneo, posicion.")
set_font(r2, size=11)

add_numbered("Para cada campo de categoría, crea un diccionario de mapeo. Ejemplo:")
add_code_inline("mapeo_generos = {")
add_code_inline("    \"rock\":         \"Rock\",")
add_code_inline("    \"rokc\":         \"Rock\",   # typo")
add_code_inline("    \"electro\":      \"Electrónica\",")
add_code_inline("    \"electronica\":  \"Electrónica\",")
add_code_inline("    ...}")

add_numbered("Crea una función normalizar_categoria(valor, diccionario_mapeo) que:")
add_bullet("Busque el valor (en minúsculas y sin espacios extra) en el diccionario.", level=1)
add_bullet("Devuelva el valor normalizado si lo encuentra.", level=1)
add_bullet("Si no encuentra coincidencia, devuelva el valor original e imprima un aviso:", level=1)
add_code_inline("    AVISO: valor no reconocido en genero_musical → 'folk-electrónico'")

add_numbered("Aplica la normalización a todos los campos de categoría.")
add_numbered("Imprime un resumen de valores únicos ANTES y DESPUÉS de normalizar:")
add_code_inline("genero_musical ANTES: Rock, rock, ROCK, Rokc, Electrónica, electro, electronica...")
add_code_inline("genero_musical DESPUÉS: Rock, Electrónica, Pop, Jazz...")

add_pista([
    "Pistas:",
    "Para construir los diccionarios de mapeo, primero extrae todos los valores únicos de ese campo",
    "  con un set(): valores_unicos = set(r['genero_musical'].lower().strip() for r in datos)",
    "  Imprime el set, observa las variaciones y construye el mapeo.",
    "Busca siempre en minúsculas: diccionario.get(valor.lower().strip(), valor)",
])

add_entregable("Datos con todas las categorías unificadas.")

# ─── FASE 6 ───────────────────────────────────────────────────────────────────
add_heading_fase(6, "NORMALIZACIÓN DE FECHAS", "⏱ ~2 horas")

add_heading3("Objetivo")
add_body("Unificar todos los formatos de fecha a un único formato estándar: DD/MM/AAAA.")

add_heading3("Tareas")
add_numbered("El formato estándar elegido es: DD/MM/AAAA  (ejemplo: 15/07/2026).")
add_numbered("Identifica todos los formatos de fecha que aparecen en los datos.")
add_numbered("Crea una función normalizar_fecha(fecha_texto) que detecte el formato y convierta la fecha al estándar.")
add_numbered("Aplica la función a todos los campos de fecha de los 4 ficheros.")
add_numbered("Imprime cuántas fechas se han normalizado y cuántas no se han podido convertir.")

add_heading3("Formatos que encontrarás en los datos")
add_table_simple(
    ["Formato", "Cómo detectarlo", "Ejemplo"],
    [
        ["DD/MM/AAAA",       "Contiene '/' y el primer trozo tiene 1-2 dígitos",        "15/07/2026"],
        ["AAAA-MM-DD",       "Contiene '-' y el primer trozo tiene 4 dígitos",          "2026-07-15"],
        ["DD de mes de AAAA","Contiene ' de '",                                          "15 de julio de 2026"],
        ["mes DD, AAAA",     "Contiene ',' y el primer trozo es texto",                  "julio 15, 2026"],
        ["DD-MM-AAAA",       "Contiene '-' y el primer trozo tiene 1-2 dígitos",        "15-07-2026"],
        ["DD/M/AA",          "Contiene '/' y el año tiene 2 dígitos",                   "15/7/26"],
        ["DD-mes-AAAA",      "Contiene '-' y el trozo del medio es texto",              "15-jul-2026"],
    ],
    col_widths=[4, 6, 4]
)

add_heading3("Diccionario de meses (para convertir texto a número)")
add_code_inline("meses = {")
add_code_inline("    \"enero\":1, \"febrero\":2, \"marzo\":3, \"abril\":4,")
add_code_inline("    \"mayo\":5, \"junio\":6, \"julio\":7, \"agosto\":8,")
add_code_inline("    \"septiembre\":9, \"octubre\":10, \"noviembre\":11, \"diciembre\":12,")
add_code_inline("    \"jan\":1, \"feb\":2, \"mar\":3, \"apr\":4, \"may\":5, \"jun\":6,")
add_code_inline("    \"jul\":7, \"aug\":8, \"sep\":9, \"oct\":10, \"nov\":11, \"dec\":12")
add_code_inline("}")

add_pista([
    "Pistas:",
    "Usa if/elif para detectar el formato según su estructura antes de parsear.",
    "Para años de 2 dígitos (ej: '26'), añade 2000: anio = int(anio) + 2000",
    "Para el formato 'DD de mes de AAAA': partes = fecha.split(' de ') → partes[0]=día, partes[1]=mes, partes[2]=año",
    "Para el formato 'mes DD, AAAA': partes = fecha.replace(',','').split() → partes[0]=mes, partes[1]=día, partes[2]=año",
    "Si ningún formato coincide, devuelve 'FECHA INVÁLIDA' e imprímelo como aviso.",
    "Al final, formatea siempre así: f\"{dia:02d}/{mes:02d}/{anio}\"",
])

add_entregable("Datos con todas las fechas en formato DD/MM/AAAA.")

# ─── FASE 7 ───────────────────────────────────────────────────────────────────
add_heading_fase(7, "ELIMINACIÓN DE DUPLICADOS", "⏱ ~2 horas")

add_heading3("Objetivo")
add_body("Detectar y eliminar registros duplicados en cada fichero, conservando siempre el registro más completo.")

add_heading3("¿Qué hace que dos registros sean \"el mismo\"?")

p = doc.add_paragraph()
r = p.add_run("Opción A: ")
set_font(r, size=11, bold=True)
add_bullet("artistas: mismo nombre (ignorando mayúsculas y espacios)", level=0)
add_bullet("ventas: mismo id_venta  O  (mismo nombre + misma fecha + mismo tipo de entrada)", level=0)
add_bullet("patrocinadores: mismo nombre_empresa", level=0)
add_bullet("horarios: mismo escenario + misma fecha + misma hora_inicio", level=0)

p = doc.add_paragraph()
r = p.add_run("Opción B: ")
set_font(r, size=11, bold=True)
add_bullet("jugadores: mismo gamertag (ignorando mayúsculas y espacios)", level=0)
add_bullet("partidas: mismo id_partida  O  (mismos equipos + misma fecha)", level=0)
add_bullet("premios: mismo torneo + misma edición + mismo equipo", level=0)
add_bullet("staff: mismo nombre + mismo equipo", level=0)

add_heading3("Tareas")
add_numbered("Crea una función eliminar_duplicados(datos, campos_clave) que:")
add_bullet("Reciba la lista de diccionarios y los campos que identifican un registro único.", level=1)
add_bullet("Compare registros normalizando los campos clave (minúsculas + sin espacios extra).", level=1)
add_bullet("Cuando encuentre duplicados, conserve el registro más completo (el que tenga menos campos vacíos o None).", level=1)
add_bullet("Devuelva la lista sin duplicados.", level=1)

add_numbered("Aplica la función a los 4 ficheros.")
add_numbered("Imprime cuántos duplicados se han eliminado:")
add_code_inline("artistas.csv: 15 duplicados eliminados (250 → 235 registros)")

add_pista([
    "Pistas:",
    "Para construir la clave de un registro: clave = tuple(str(r[c]).lower().strip() for c in campos_clave)",
    "Usa un diccionario auxiliar {clave: registro} para detectar duplicados.",
    "Para comparar cuál es más completo: cuenta cuántos campos tienen valor None o vacío.",
    "  El registro con MENOS campos vacíos es el que conservas.",
    "  Ejemplo: sum(1 for v in r.values() if not v or v == 'SIN DATOS')",
])

add_entregable("Datos sin duplicados, con el conteo de registros eliminados.")

# ─── FASE 8 ───────────────────────────────────────────────────────────────────
add_heading_fase(8, "DETECCIÓN DE VALORES FUERA DE RANGO", "⏱ ~1.5 horas")

add_heading3("Objetivo")
add_body("Detectar y corregir valores numéricos que no tienen sentido para el contexto.")

add_heading3("Rangos válidos por opción")

p = doc.add_paragraph()
r = p.add_run("Opción A:")
set_font(r, size=11, bold=True)
add_table_simple(
    ["Campo", "Mínimo", "Máximo", "Qué hacer si está fuera de rango"],
    [
        ["cache_eur",          "500",    "500.000", "Negativo → valor absoluto. Absurdo → \"REVISAR MANUALMENTE\""],
        ["precio (entrada)",   "20",     "500",     "Negativo → valor absoluto. Absurdo → \"REVISAR MANUALMENTE\""],
        ["importe_patrocinio", "1.000",  "1.000.000","Negativo → valor absoluto. Absurdo → \"REVISAR MANUALMENTE\""],
    ],
    col_widths=[4, 2, 2.5, 7]
)

p = doc.add_paragraph()
r = p.add_run("Opción B:")
set_font(r, size=11, bold=True)
add_table_simple(
    ["Campo", "Mínimo", "Máximo", "Qué hacer si está fuera de rango"],
    [
        ["edad",               "14",  "45",      "Fuera de rango → \"REVISAR MANUALMENTE\""],
        ["salario_mensual",    "500", "50.000",  "Negativo → valor absoluto. Absurdo → \"REVISAR MANUALMENTE\""],
        ["puntuacion",         "0",   "30",      "Negativo → 0. Superior al máximo → \"REVISAR MANUALMENTE\""],
        ["duracion_minutos",   "5",   "120",     "Fuera de rango → \"REVISAR MANUALMENTE\""],
        ["premio_eur",         "100", "1.000.000","Negativo → valor absoluto. Absurdo → \"REVISAR MANUALMENTE\""],
    ],
    col_widths=[4, 2, 2.5, 7]
)

add_heading3("Tareas")
add_numbered("Crea una función detectar_fuera_de_rango(valor, minimo, maximo) que devuelva True si el valor está fuera de rango.")
add_numbered("Para cada valor fuera de rango, aplica la acción indicada en la tabla.")
add_numbered("Imprime un listado de todos los valores fuera de rango con el registro completo:")
add_code_inline("FUERA DE RANGO en artistas.csv, registro 47: cache_eur = -500  → corregido a 500")
add_code_inline("FUERA DE RANGO en artistas.csv, registro 12: cache_eur = 9999999 → REVISAR MANUALMENTE")

add_pista([
    "Pistas:",
    "Recuerda hacer la conversión numérica (Fase 4) ANTES de aplicar esta fase.",
    "Para el valor absoluto usa abs(valor).",
    "Los valores marcados como 'REVISAR MANUALMENTE' deben aparecer también en el informe final.",
])

add_entregable("Datos con valores numéricos dentro de rangos lógicos y lista de casos a revisar.")

# ─── FASE 9 ───────────────────────────────────────────────────────────────────
add_heading_fase(9, "VALIDACIÓN CRUZADA ENTRE FICHEROS", "⏱ ~1.5 horas")

add_heading3("Objetivo")
add_body(
    "Verificar que los datos son coherentes entre ficheros. No basta con que cada fichero "
    "esté limpio por separado: los valores que aparecen en un fichero deben existir en el "
    "fichero de referencia correspondiente."
)

add_heading3("Relaciones entre ficheros a validar")

p = doc.add_paragraph()
r = p.add_run("Opción A:")
set_font(r, size=11, bold=True)
add_bullet("Los artistas en escenarios_horarios deben existir en artistas.csv.", level=0)
add_bullet("Los nombres de escenario deben ser consistentes entre ficheros.", level=0)
add_bullet("Las fechas de venta de entradas deben ser anteriores al inicio del festival.", level=0)

p = doc.add_paragraph()
r = p.add_run("Opción B:")
set_font(r, size=11, bold=True)
add_bullet("Los equipos en jugadores deben existir en equipos.xlsx.", level=0)
add_bullet("Los equipos en partidas deben existir en equipos.xlsx.", level=0)
add_bullet("Los equipos en premios deben existir en equipos.xlsx.", level=0)
add_bullet("Los torneos en premios deben aparecer también en partidas.json.", level=0)

add_heading3("Cómo resolver las inconsistencias")
add_body("Para cada valor que no encuentres en el fichero de referencia, sigue este proceso:")
add_numbered("Normaliza el valor (lower + strip) e inténtalo de nuevo.")
add_numbered("Si tras normalizar sí hay coincidencia → corrige el valor al nombre oficial del fichero maestro.")
add_numbered("Si tras normalizar sigue sin haber coincidencia → NO borres el registro. En su lugar:")
add_bullet("Deja el valor original intacto.", level=1)
add_bullet("Añade una columna nueva llamada estado_referencia al registro.", level=1)
add_bullet("Ponla a \"OK\" si la referencia se encontró, o \"REVISAR\" si no se encontró.", level=1)
add_bullet("Añade el registro al informe de avisos del fichero de salida.", level=1)

add_numbered("Vuelve a ejecutar la validación para confirmar que ya no hay nuevas inconsistencias resoluble automáticamente.")

add_code_inline("Ejemplo de registro con estado_referencia:")
add_code_inline("{")
add_code_inline("  \"equipo\": \"Arctic Fxoes\",   ← valor original conservado")
add_code_inline("  \"puntuacion\": 8,")
add_code_inline("  ...,")
add_code_inline("  \"estado_referencia\": \"REVISAR\"  ← columna nueva añadida")
add_code_inline("}")

add_pista([
    "Pistas:",
    "Construye un set con los nombres normalizados del fichero maestro para hacer búsquedas rápidas:",
    "  equipos_validos = set(e['nombre_equipo'].lower().strip() for e in equipos)",
    "Luego para cada registro: if equipo.lower().strip() in equipos_validos → OK, si no → REVISAR",
    "Para añadir la columna: simplemente haz registro['estado_referencia'] = 'OK' (o 'REVISAR').",
])

add_entregable("Datos con la columna estado_referencia añadida e informe de inconsistencias.")

# ─── FASE 10 ──────────────────────────────────────────────────────────────────
add_heading_fase(10, "EXPORTACIÓN Y REPORTE FINAL", "⏱ ~1.5 horas")

add_heading3("Objetivo")
add_body("Generar los ficheros limpios finales y un informe resumen de todo el proceso.")

add_heading3("Tareas")
add_numbered("Crea la carpeta datos_limpios/ dentro de tu carpeta de opción (usa os.makedirs).")
add_numbered("Exporta los datos limpios:")
add_bullet("Un fichero CSV por cada fichero original (usa csv.DictWriter).", level=1)
add_bullet("Un fichero Excel (datos_completos.xlsx) con todos los datos en hojas separadas, una por fichero.", level=1)

add_numbered("Crea el fichero informe_limpieza.txt con este contenido:")
add_code_inline("=== INFORME DE LIMPIEZA ===")
add_code_inline("Fecha del proceso: DD/MM/AAAA")
add_code_inline("Ficheros procesados: 4")
add_code_inline("")
add_code_inline("--- RESUMEN POR FICHERO ---")
add_code_inline("artistas.csv:")
add_code_inline("  Registros originales: 250 | Registros finales: 231")
add_code_inline("  Duplicados eliminados: 15")
add_code_inline("  Valores vacíos tratados: 47")
add_code_inline("  Categorías normalizadas: 89")
add_code_inline("  Fechas convertidas: 0  (sin fechas en este fichero)")
add_code_inline("  Valores fuera de rango corregidos: 5")
add_code_inline("")
add_code_inline("--- VALIDACIÓN CRUZADA ---")
add_code_inline("  Inconsistencias resueltas automáticamente: 12")
add_code_inline("  Registros marcados como REVISAR: 3")
add_code_inline("")
add_code_inline("--- AVISOS (requieren atención humana) ---")
add_code_inline("  REVISAR MANUALMENTE (fuera de rango): 2 casos")
add_code_inline("  REVISAR (sin referencia en fichero maestro): 3 casos")
add_code_inline("    · partidas.json, registro 47: equipo = 'Arctic Fxoes'")

add_numbered("Imprime el informe por pantalla Y guárdalo en el fichero.")

add_pista([
    "Pistas:",
    "Para escribir el informe en fichero: abre con open('datos_limpios/informe_limpieza.txt', 'w', encoding='utf-8')",
    "  y usa fichero.write(linea + '\\n') para cada línea.",
    "Para el Excel con varias hojas: crea un Workbook, y para cada conjunto de datos haz",
    "  wb.create_sheet(title='nombre_hoja') antes de escribir los datos.",
    "Para la fecha del proceso usa: from datetime import date  →  date.today().strftime('%d/%m/%Y')",
])

add_entregable("Carpeta datos_limpios/ con los CSV limpios, el Excel consolidado y el informe.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CRITERIOS DE EVALUACIÓN
# ══════════════════════════════════════════════════════════════════════════════

add_heading1("CRITERIOS DE EVALUACIÓN")

add_table_simple(
    ["Criterio", "Peso", "Qué se evalúa"],
    [
        ["Funcionalidad", "40%",
         "¿El código carga los 4 ficheros correctamente? ¿Detecta los problemas? ¿Limpia y normaliza bien? ¿Los ficheros de salida son correctos?"],
        ["Código",        "30%",
         "¿Está organizado en funciones reutilizables? ¿Las funciones reciben parámetros y devuelven valores? ¿Se manejan excepciones donde hace falta? ¿Es legible?"],
        ["Proceso",       "20%",
         "¿Se ha seguido el orden de las fases? ¿Cada fase produce un resultado visible? ¿El informe final documenta el proceso?"],
        ["Rigor",         "10%",
         "¿Se ha validado que no quedan problemas sin tratar? ¿Los datos cruzan correctamente? ¿Se han documentado las decisiones?"],
    ],
    col_widths=[4, 2, 10]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  CONSEJOS
# ══════════════════════════════════════════════════════════════════════════════

add_heading1("CONSEJOS ANTES DE EMPEZAR")

consejos = [
    ("Lee los datos antes de escribir código.",
     "Abre los ficheros, mira qué contienen, entiende la estructura. No te lances a programar hasta saber exactamente con qué estás trabajando."),
    ("Imprime mucho.",
     "Cada vez que hagas una transformación, imprime un antes y después para verificar que funciona. No asumas que tu código hace lo correcto: compruébalo."),
    ("Usa funciones.",
     "No escribas todo en un bloque. Crea funciones pequeñas que hagan una cosa bien y combínalas. El ej10 del curso es un buen modelo a seguir."),
    ("Prueba con pocos datos primero.",
     "Si un fichero tiene 400 registros, prueba tu función con los 5 primeros antes de aplicarla a todos."),
    ("No modifiques los ficheros originales.",
     "Trabaja siempre con copias en memoria. Los ficheros de la carpeta datos/ deben quedar intactos."),
    ("Comenta las decisiones.",
     "Cuando decidas cómo tratar un valor vacío o un outlier, escribe un comentario explicando por qué. Eso forma parte de la evaluación."),
    ("Pide ayuda si llevas más de 20-30 minutos atascado.",
     "Pero primero intenta buscar la solución por tu cuenta. El error de Python que ves en pantalla casi siempre te dice exactamente qué está mal."),
]

for titulo, desc in consejos:
    p = doc.add_paragraph()
    r1 = p.add_run(f"▸ {titulo}  ")
    set_font(r1, size=11, bold=True, color=AZUL_OSCURO)
    r2 = p.add_run(desc)
    set_font(r2, size=11)
    p.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════════════════════
#  GUARDAR
# ══════════════════════════════════════════════════════════════════════════════

ruta = "/Volumes/Proyectos/proyectos/practica_final/enunciado_practica_final.docx"
doc.save(ruta)
print(f"✅ Documento generado: {ruta}")
